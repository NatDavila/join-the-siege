import io
import pytest
from PIL import Image
import docx
import openpyxl
import xlwt
from src.utils import (
    preprocess_image,
    ocr_image,
    ocr_pdf,
    ocr_docx,
    ocr_excel,
    text_extractor,
)

SAMPLE_FILE_DIR = "test_sample_files"


def test_preprocess_image():
    """
    Loads an image from test sample files and tests grayscale preprocessing
    """
    image_path = f"{SAMPLE_FILE_DIR}/test_image_ocr.png"
    
    img = Image.open(image_path)  
    processed = preprocess_image(img)

    assert processed.mode == "L"


def test_ocr_image():
    """
    Loads an image from test sample files and tests OCR extraction
    """
    image_path = f"{SAMPLE_FILE_DIR}/test_image_ocr.png"
    
    with open(image_path, "rb") as img_file:
        image_bytes = img_file.read()

    assert ocr_image(image_bytes) == "Nico rocks\n"


def test_ocr_pdf():
    """
    Loads a PDF with text from test sample files and tests OCR extraction
    """
    pdf_path = f"{SAMPLE_FILE_DIR}/test_pdf_ocr.pdf"

    with open(pdf_path, "rb") as pdf_file:
        pdf_bytes = pdf_file.read()

    print(ocr_pdf(pdf_bytes))
    assert ocr_pdf(pdf_bytes) == "i should start up a Minecraft server again\n"


def test_ocr_docx():
    """
    Creates a simple DOCX file and tests OCR extraction
    """
    doc = docx.Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test.")
    buf = io.BytesIO()
    doc.save(buf)

    assert ocr_docx(buf.getvalue()) == "Hello World\nThis is a test.\n"


def test_ocr_xlsx(ext, expected_texts):
    """
    Creates a simple XLSX file and tests OCR extraction
    """
    ext = "xlsx"
    expected_texts = ["Data1", "Data2"]
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = expected_texts[0]
    ws["B1"] = expected_texts[1]
    buf = io.BytesIO()
    wb.save(buf)
    excel_bytes = buf.getvalue()

    result = ocr_excel(excel_bytes, ext)

    for text in expected_texts:
        assert text in result


def test_ocr_xlsx():
    """
    Creates a XLS file and tests OCR extraction
    """
    ext = "xls"
    expected_texts = ["Cell1", "Cell2"]
    workbook = xlwt.Workbook()
    sheet = workbook.add_sheet("Sheet1")
    sheet.write(0, 0, expected_texts[0])
    sheet.write(0, 1, expected_texts[1])
    buf = io.BytesIO()
    workbook.save(buf)
    excel_bytes = buf.getvalue()

    result = ocr_excel(excel_bytes, ext)

    for text in expected_texts:
        assert text in result


def test_text_extractor_txt():
    """
    Creates a txt file and tests text extraction
    """
    sample = "This is a text file."
    result = text_extractor(sample.encode("utf-8"), "txt")
    assert result == sample


def test_text_extractor_unsupported():
    """
    Tests text extraction from an unsupported file format
    """
    with pytest.raises(ValueError, match="Unsupported file format"):
        text_extractor(b"dummy data", "unsupported_extension")
